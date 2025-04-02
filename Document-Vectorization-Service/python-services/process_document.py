import os
import json
import xml.etree.ElementTree as ET
from dotenv import load_dotenv
from semantic_chunking import chunk_document
from embedding_function import GeminiEmbeddingFunction
from unstructured.partition.auto import partition
import docx
import asyncio
import hashlib
from urllib.parse import urlparse
from web_content_fetcher import fetch_web_content, process_batch_urls
from qdrant_wrapper import get_qdrant_client

def extract_text_from_json(file_path):
    """Extract text from JSON file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return json.dumps(data, indent=2)

def extract_text_from_xml(file_path):
    """Extract text from XML file."""
    tree = ET.parse(file_path)
    root = tree.getroot()
    
    def extract_text(element):
        text = element.text.strip() if element.text else ""
        for child in element:
            text += " " + extract_text(child)
        return text.strip()
    
    return extract_text(root)

def extract_text_from_markdown(file_path):
    """Extract text from Markdown file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    # Convert markdown to plain text (simple approach)
    return text

def extract_text_from_txt(file_path):
    """Extract text from plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        print(f"Text file content length: {len(text)}")
        
        # If the text is empty, add a placeholder
        if not text or len(text.strip()) == 0:
            text = "This is an empty text file."
            print("Text file was empty, adding placeholder text")
        
        return text
    except UnicodeDecodeError:
        # Try with different encoding if utf-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            print(f"Successfully read file with latin-1 encoding, length: {len(text)}")
            return text
        except Exception as e:
            print(f"Error reading text file with latin-1 encoding: {str(e)}")
            raise
    except Exception as e:
        print(f"Error reading text file: {str(e)}")
        raise

def get_database_path():
    """Get the database path."""
    base_path = os.path.join(os.path.dirname(__file__), '..', 'vector-database')
    db_path = os.path.join(base_path, "store-new")
    os.makedirs(db_path, exist_ok=True)
    return db_path

def process_document(file_path, collection_name="default", metadata=None):
    """Process a document and store it in Qdrant.
    
    Args:
        file_path (str): Path to the document to process
        collection_name (str): Name of the collection to store embeddings in.
        metadata (dict, optional): Additional metadata to store with the document chunks.
    """
    print(f"🚀 Processing document: {file_path}")
    
    try:
        # Load environment variables
        load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))
        api_key = os.getenv('GOOGLE_GEMINI_API_KEY')
        
        if not api_key:
            raise ValueError("GOOGLE_GEMINI_API_KEY environment variable not set")
        
        # Extract text from document
        print("🔹 Extracting text...")
        extension = os.path.splitext(file_path)[1].lower()
        
        # Print file info for debugging
        print(f"File size: {os.path.getsize(file_path)} bytes")
        print(f"File extension: {extension}")
        
        if extension == '.json':
            text = extract_text_from_json(file_path)
        elif extension == '.xml':
            text = extract_text_from_xml(file_path)
        elif extension == '.md':
            text = extract_text_from_markdown(file_path)
        elif extension == '.txt':
            text = extract_text_from_txt(file_path)
        else:
            try:
                elements = partition(filename=file_path)
                text = "\n".join([str(el) for el in elements])
            except Exception as e:
                print(f"Error using unstructured partition: {str(e)}")
                # Fallback to simple text extraction for unknown file types
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                except UnicodeDecodeError:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        text = f.read()
        
        if not text:
            # If no text was extracted, add a placeholder
            print("Warning: No text extracted, using placeholder")
            text = f"This document ({os.path.basename(file_path)}) appears to be empty or could not be processed."
        
        print(f"Extracted text length: {len(text)}")
        print("✅ Text extracted successfully")
        
        # Chunk text using semantic chunking
        print("🔹 Applying semantic chunking...")
        document_name = os.path.basename(file_path)
        chunks, chunk_metadatas = chunk_document(
            text=text,
            document_name=document_name,
            metadata=metadata,
            chunk_size=4000,  # Larger chunks for better context
            chunk_overlap=200  # Overlap to maintain context between chunks
        )
        print(f"✅ Created {len(chunks)} semantic chunks")
        
        # Initialize Qdrant client
        print("🔹 Creating embeddings and storing in Qdrant...")
        qdrant_client = get_qdrant_client()
        
        embedding_model = GeminiEmbeddingFunction(api_key=api_key)
        collection = qdrant_client.get_or_create_collection(
            name=collection_name,
            embedding_function=embedding_model
        )
        
        # Generate unique document IDs by using the filename
        file_identifier = os.path.basename(file_path).replace(".", "_")  # Unique filename-based ID
        new_doc_ids = [f"{file_identifier}_doc_{i}" for i in range(len(chunks))]
        
        collection.add(
            documents=chunks,
            ids=new_doc_ids,
            metadatas=chunk_metadatas
        )
        
        print(f"✅ Successfully added {len(new_doc_ids)} new chunks to Qdrant")
        print("✅ Document processing complete!")
        
    except Exception as e:
        print(f"❌ Error processing document: {str(e)}")
        raise

async def process_url(url, collection_name="default", metadata=None, follow_links=True, max_links=5):
    """
    Process a URL, fetch its content and store it in Qdrant.
    
    Args:
        url (str): URL to process
        collection_name (str): Name of the collection to store the document in
        metadata (dict, optional): Additional metadata to store with the document
        follow_links (bool): Whether to follow links within the page (one level deep)
        max_links (int): Maximum number of links to follow
        
    Returns:
        dict: Information about the processed document
    """
    print(f"🚀 Processing URL: {url}")
    
    try:
        # Load environment variables
        load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))
        api_key = os.getenv('GOOGLE_GEMINI_API_KEY')
        
        if not api_key:
            raise ValueError("GOOGLE_GEMINI_API_KEY environment variable not set")
        
        # Fetch and process web content
        print("🔹 Fetching and extracting content...")
        from web_content_fetcher import fetch_web_content
        
        # Fetch content with link following if enabled
        content, content_metadata, content_type = fetch_web_content(url, follow_links=follow_links, max_links=max_links)
        
        if not content:
            print("Warning: No content extracted, using placeholder")
            content = f"This URL ({url}) appears to be empty or could not be processed."
        
        print(f"Extracted content length: {len(content)}")
        print(f"Content type: {content_type}")
        print("✅ Content extracted successfully")
        
        # Chunk text
        print("🔹 Chunking content...")
        chunks = chunk_document(content)
        print(f"✅ Created {len(chunks)} chunks")
        
        # Initialize Qdrant
        print("🔹 Creating embeddings and storing in Qdrant...")
        qdrant_client = get_qdrant_client()
        
        embedding_model = GeminiEmbeddingFunction(api_key=api_key)
        collection = qdrant_client.get_or_create_collection(
            name=collection_name,
            embedding_function=embedding_model
        )
        
        # Generate unique document IDs by hashing the URL
        url_hash = hashlib.md5(url.encode()).hexdigest()
        new_doc_ids = [f"{url_hash}_doc_{i}" for i in range(len(chunks))]
        
        # Prepare metadata for each chunk
        if metadata is None:
            metadata = {}
        
        # Merge content metadata with provided metadata
        merged_metadata = {**content_metadata, **metadata}
        
        # Add document name to metadata if not present
        if "document_name" not in merged_metadata:
            # Use the title from content metadata or fallback to domain name
            document_name = content_metadata.get("title", urlparse(url).netloc)
            merged_metadata["document_name"] = document_name
            
        # Add URL as source in metadata
        merged_metadata["source"] = url
        merged_metadata["source_type"] = "web"
        
        # Create metadata list for each chunk
        metadatas = [merged_metadata.copy() for _ in range(len(chunks))]
        
        collection.add(
            documents=chunks,
            ids=new_doc_ids,
            metadatas=metadatas
        )
        
        print(f"✅ Successfully added {len(new_doc_ids)} new chunks to Qdrant")
        print("✅ URL processing complete!")
        
        return {
            "status": "success",
            "document_id": url_hash,
            "url": url,
            "content_type": content_type,
            "collection": collection_name,
            "chunks": len(chunks)
        }
        
    except Exception as e:
        error_msg = f"Error processing URL {url}: {str(e)}"
        print(error_msg)
        return {
            "status": "error",
            "url": url,
            "error": str(e)
        }

async def process_urls_batch(urls, collection_name="default", metadata=None, follow_links=True, max_links=3):
    """
    Process multiple URLs in batch.
    
    Args:
        urls (list): List of URLs to process
        collection_name (str): Name of the collection to store the documents in
        metadata (dict, optional): Additional metadata to store with the documents
        follow_links (bool): Whether to follow links within each page (one level deep)
        max_links (int): Maximum number of links to follow per URL
        
    Returns:
        list: Results for each URL
    """
    results = []
    
    for url in urls:
        try:
            result = await process_url(url, collection_name, metadata, follow_links, max_links)
            results.append(result)
        except Exception as e:
            print(f"Error in batch processing for URL {url}: {str(e)}")
            results.append({
                "status": "error",
                "url": url,
                "error": str(e)
            })
    
    return results

def extract_text(file_path):
    """Extract text from a document based on file type."""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.json':
            return extract_text_from_json(file_path)
        elif file_extension == '.xml':
            return extract_text_from_xml(file_path)
        elif file_extension == '.md':
            return extract_text_from_markdown(file_path)
        elif file_extension == '.txt':
            return extract_text_from_txt(file_path)
        elif file_extension == '.docx':
            # Special handling for DOCX files with fallback
            try:
                # First try with unstructured
                elements = partition(filename=file_path)
                return "\n\n".join([str(element) for element in elements])
            except Exception as docx_error:
                print(f"Error using unstructured for DOCX: {str(docx_error)}")
                print("Attempting to use python-docx as fallback...")
                
                # Try to import python-docx
                try:
                    import docx
                    doc = docx.Document(file_path)
                    full_text = []
                    
                    # Extract text from paragraphs
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                row_text.append(cell.text)
                            full_text.append(" | ".join(row_text))
                    
                    return "\n\n".join(full_text)
                except ImportError:
                    print("python-docx not available, trying with basic file reading...")
                    # Last resort - try to read as binary and extract text
                    try:
                        with open(file_path, 'rb') as f:
                            content = f.read()
                        # Extract any readable text from binary content
                        text_content = ""
                        for i in range(0, len(content), 2):
                            if i + 1 < len(content):
                                char_code = content[i] + (content[i+1] << 8)
                                if 32 <= char_code < 127:  # ASCII printable chars
                                    text_content += chr(char_code)
                                elif char_code == 13 or char_code == 10:  # CR/LF
                                    text_content += "\n"
                        return text_content
                    except Exception as basic_error:
                        print(f"Basic text extraction failed: {str(basic_error)}")
                        return "This DOCX file could not be processed due to missing dependencies. Please install unstructured[docx] or python-docx."
        else:
            # Use unstructured for other file types
            try:
                elements = partition(filename=file_path)
                return "\n\n".join([str(element) for element in elements])
            except Exception as partition_error:
                print(f"Error with partition for {file_extension}: {str(partition_error)}")
                # Try to read file as text as fallback
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except:
                    try:
                        with open(file_path, 'r', encoding='latin-1') as f:
                            return f.read()
                    except Exception as read_error:
                        print(f"Fallback text reading failed: {str(read_error)}")
                        return f"This file with extension {file_extension} could not be processed."
    except Exception as e:
        print(f"❌ Error extracting text: {str(e)}")
        return None

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python process_document.py <file_path> [collection_name]")
        sys.exit(1)
    
    file_path = sys.argv[1]
    collection_name = sys.argv[2] if len(sys.argv) > 2 else "default"
    
    process_document(file_path, collection_name)
